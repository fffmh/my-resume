# -*- coding: utf-8 -*-
"""从本地文件抽取纯文本：txt/md 直读、docx 用 python-docx、pdf 用 PyMuPDF。"""
import io

import docx
import fitz


def extract_text(filename: str, data: bytes):
    name = filename.lower()
    try:
        if name.endswith(('.txt', '.md', '.csv', '.log')):
            return data.decode('utf-8', errors='replace'), ''
        if name.endswith('.docx'):
            doc = docx.Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(' | '.join(cells))
            return '\n'.join(parts), ''
        if name.endswith('.pdf'):
            pdf = fitz.open(stream=data, filetype='pdf')
            text = '\n'.join(page.get_text() for page in pdf)
            pdf.close()
            return text, ''
        return '', '暂不支持该文件类型，请使用 .txt / .md / .docx / .pdf'
    except Exception as exc:  # noqa: BLE001
        return '', f'解析失败：{exc}'