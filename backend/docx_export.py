# -*- coding: utf-8 -*-
"""内置模板 Word 导出：用 python-docx 生成一份干净专业的简历 .docx（无需上传模板）。"""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x6B, 0x72, 0x80)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
FONT = '微软雅黑'


def _set_font(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)


def _text(doc, value, size=10.5, bold=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if value:
        _set_font(p.add_run(str(value)), size=size, bold=bold, color=color)
    return p


def _section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    _set_font(p.add_run(title), size=13, bold=True, color=ACCENT)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pPr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): '1F3A5F'})
    pBdr.append(bottom)
    pPr.append(pBdr)


def _bullets(doc, text):
    for line in str(text).split('\n'):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        _set_font(p.add_run('•  '), size=10.5, color=ACCENT)
        _set_font(p.add_run(line), size=10.5, color=BLACK)


def _row(doc, title, subtitle='', times=''):
    if title:
        _text(doc, title, size=11.5, bold=True, color=BLACK, space_after=1)
    meta = ' · '.join([s for s in (subtitle, times) if s])
    if meta:
        _text(doc, meta, size=9.5, color=MUTED, space_after=2)


def build_docx(data):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    basic = data.get('basic') or {}
    intention = data.get('intention') or {}
    name = str(basic.get('name', '') or '未填写姓名')
    role = str(data.get('targetJob', '') or intention.get('position', '') or '')
    _text(doc, name, size=22, bold=True, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    if role:
        _text(doc, role, size=12, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    contact = ' | '.join([str(basic[k]) for k in ('phone', 'email', 'city') if str(basic.get(k, '')).strip()])
    if basic.get('years'):
        contact = (contact + ' | ' + str(basic['years'])) if contact else str(basic['years'])
    if contact:
        _text(doc, contact, size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    if intention.get('position'):
        _section_title(doc, '求职意向')
        _text(doc, ' · '.join([str(intention[k]) for k in ('position', 'salary', 'city', 'join') if str(intention.get(k, '')).strip()]), space_after=4)

    def entries(sid):
        return data.get(sid) or []

    if entries('work'):
        _section_title(doc, '工作经历')
        for w in entries('work'):
            _row(doc, ' · '.join([str(w.get('position', '')), str(w.get('company', ''))]),
                 times=f"{w.get('start', '')} - {w.get('end', '')}")
            _bullets(doc, w.get('content', ''))
            if w.get('achievement'):
                _bullets(doc, w.get('achievement', ''))

    if entries('project'):
        _section_title(doc, '项目经历')
        for p in entries('project'):
            _row(doc, ' · '.join([str(p.get('name', '')), str(p.get('role', ''))]),
                 subtitle=str(p.get('tech', '')))
            _bullets(doc, p.get('desc', ''))
            if p.get('contribution'):
                _bullets(doc, p.get('contribution', ''))

    if entries('education'):
        _section_title(doc, '教育经历')
        for e in entries('education'):
            _row(doc, ' · '.join([str(e.get('school', '')), str(e.get('major', ''))]),
                 times=f"{e.get('start', '')} - {e.get('end', '')}")
            if e.get('honor'):
                _bullets(doc, e.get('honor', ''))

    if entries('skills'):
        _section_title(doc, '技能特长')
        _text(doc, '、'.join([str(s.get('name', '')) for s in entries('skills') if str(s.get('name', '')).strip()]), space_after=4)

    if entries('certificate'):
        _section_title(doc, '证书资质')
        _text(doc, '、'.join([str(c.get('name', '')) for c in entries('certificate') if str(c.get('name', '')).strip()]), space_after=4)

    if data.get('self'):
        _section_title(doc, '自我评价')
        _text(doc, str(data['self']), space_after=4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()