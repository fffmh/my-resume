# -*- coding: utf-8 -*-
import io

import docx
import fitz

from backend import fill


def _make_docx_template():
    d = docx.Document()
    d.add_paragraph('姓名：{{姓名}}')
    d.add_paragraph('{% for w in 工作经历 %}职位：{{ w.position }} - {{ w.company }}{% endfor %}')
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_fill_docx():
    tpl = _make_docx_template()
    data = {
        'basic': {'name': '张三'},
        'work': [{'position': '前端', 'company': 'A公司'}, {'position': '后端', 'company': 'B公司'}],
    }
    out = fill.fill_docx(tpl, data)
    doc = docx.Document(io.BytesIO(out))
    texts = [p.text for p in doc.paragraphs]
    joined = '\n'.join(texts)
    assert '张三' in joined
    assert 'A公司' in joined
    assert 'B公司' in joined
    assert '{{' not in joined
    assert '{%' not in joined


def _make_pdf_template():
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), '姓名：{{姓名}}', fontname='china-s', fontsize=12)
    buf = io.BytesIO(pdf.tobytes())
    pdf.close()
    return buf.getvalue()


def test_fill_pdf():
    tpl = _make_pdf_template()
    out = fill.fill_pdf(tpl, {'basic': {'name': '李四'}})
    doc = fitz.open(stream=out, filetype='pdf')
    try:
        text = doc[0].get_text()
    finally:
        doc.close()
    assert '李四' in text
    assert '{{' not in text


def test_fill_pdf_form_field():
    pdf = fitz.open()
    page = pdf.new_page()
    widget = fitz.Widget()
    widget.rect = fitz.Rect(72, 72, 320, 92)
    widget.field_name = '姓名'
    widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
    widget.field_value = ''
    page.add_widget(widget)
    buf = io.BytesIO(pdf.tobytes())
    pdf.close()
    out = fill.fill_pdf(buf.getvalue(), {'basic': {'name': '王五'}})
    doc = fitz.open(stream=out, filetype='pdf')
    try:
        widgets = list(doc[0].widgets())
        assert widgets and widgets[0].field_value == '王五'
    finally:
        doc.close()


def test_fill_pdf_multiline():
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), '{{自我评价}}', fontname='china-s', fontsize=12)
    buf = io.BytesIO(pdf.tobytes())
    pdf.close()
    value = '第一行：专注工程化。\n第二行：习惯用数据衡量结果。'
    out = fill.fill_pdf(buf.getvalue(), {'self': value})
    doc = fitz.open(stream=out, filetype='pdf')
    try:
        text = doc[0].get_text()
    finally:
        doc.close()
    assert '第一行' in text
    assert '第二行' in text
    assert '{{' not in text